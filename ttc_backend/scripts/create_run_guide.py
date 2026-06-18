import os
import sys
import subprocess

def install_and_import(package):
    try:
        import docx
    except ImportError:
        print(f"Installing {package}...")
        # Run pip using the current Python executable
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])
        import docx
    return docx

docx = install_and_import('python-docx')
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_guide():
    doc = Document()
    
    # Page Setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles Setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TT Cell Vocational Training Management Portal (TTC-VTP)\nRun Guide & Setup Instructions")
    run.font.name = 'Calibri'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(74, 99, 49) # #4A6331 matching theme
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run("Step-by-step instructions for running the application on another computer")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(122, 139, 153) # Grey
    
    doc.add_paragraph() # Spacer

    # Section: Prerequisites
    h1 = doc.add_paragraph()
    run_h1 = h1.add_run("1. Prerequisites & Installations")
    run_h1.font.bold = True
    run_h1.font.size = Pt(14)
    run_h1.font.color.rgb = RGBColor(74, 99, 49)
    
    doc.add_paragraph("Before running the portal on the new computer, ensure the following software is downloaded and installed:")
    
    bullets = [
        ("Python (Version 3.10 to 3.12)", "Download from https://www.python.org/downloads/. IMPORTANT: During installation, check the box that says 'Add Python to PATH'."),
        ("Node.js (LTS Version 18 or 20)", "Download from https://nodejs.org/. This installs npm (Node Package Manager) which is required for the frontend."),
        ("MongoDB Community Server", "Download from https://www.mongodb.com/try/download/community. This database runs locally on port 27017."),
        ("MongoDB Compass", "Download from https://www.mongodb.com/try/download/compass. A graphical dashboard to view database collections.")
    ]
    for title_b, desc_b in bullets:
        p = doc.add_paragraph(style='List Bullet')
        r_bold = p.add_run(f"{title_b}: ")
        r_bold.font.bold = True
        p.add_run(desc_b)

    doc.add_paragraph()

    # Section: Copying files
    h2 = doc.add_paragraph()
    run_h2 = h2.add_run("2. Copying the Project Files")
    run_h2.font.bold = True
    run_h2.font.size = Pt(14)
    run_h2.font.color.rgb = RGBColor(74, 99, 49)

    doc.add_paragraph("Transfer the project folder 'TT-CELL-VOCATIONAL-TRAINING-main/' to the destination PC.")
    
    # Note Box
    p_note = doc.add_paragraph()
    r_note_lbl = p_note.add_run("TIP: ")
    r_note_lbl.font.bold = True
    r_note_lbl.font.color.rgb = RGBColor(184, 150, 12) # Yellow/Gold
    p_note.add_run("To save time and storage space when transferring, you can delete the following auto-generated folders before copying. They will be regenerated on the new PC:")
    
    p_sub1 = doc.add_paragraph(style='List Bullet')
    p_sub1.add_run("ttcell/node_modules/").font.bold = True
    p_sub1.add_run(" (reinstalled via npm install)")
    
    p_sub2 = doc.add_paragraph(style='List Bullet')
    p_sub2.add_run("ttc_backend/venv/").font.bold = True
    p_sub2.add_run(" (recreated via python -m venv)")

    doc.add_paragraph()

    # Section: Database Check
    h3 = doc.add_paragraph()
    run_h3 = h3.add_run("3. Database Setup & Verification")
    run_h3.font.bold = True
    run_h3.font.size = Pt(14)
    run_h3.font.color.rgb = RGBColor(74, 99, 49)
    doc.add_paragraph("1. Verify that MongoDB is running on port 27017 (on Windows, it normally runs automatically as a system service).")
    doc.add_paragraph("2. Open MongoDB Compass and connect to the local URI:")
    p_uri = doc.add_paragraph()
    r_uri = p_uri.add_run("    mongodb://localhost:27017")
    r_uri.font.name = 'Consolas'
    r_uri.font.size = Pt(10.5)

    doc.add_paragraph()

    # Section: Backend Setup
    h4 = doc.add_paragraph()
    run_h4 = h4.add_run("4. Backend Setup & Run Instructions")
    run_h4.font.bold = True
    run_h4.font.size = Pt(14)
    run_h4.font.color.rgb = RGBColor(74, 99, 49)
    
    doc.add_paragraph("Open a terminal/command line window on the new PC and execute:")
    
    commands_be = [
        "cd ttc_backend",
        "# Create Python Virtual Environment",
        "python -m venv venv",
        "# Activate Virtual Environment",
        "venv\\Scripts\\activate      # Windows (CMD)",
        "& venv\\Scripts\\Activate.ps1 # Windows (PowerShell)",
        "source venv/bin/activate    # Linux / macOS",
        "# Install Dependencies",
        "pip install -r requirements.txt",
        "# Seed Database (Creates default administrator accounts)",
        "python manage.py create_admin",
        "# Start Development Server",
        "python manage.py runserver"
    ]
    for cmd in commands_be:
        p_cmd = doc.add_paragraph()
        p_cmd.paragraph_format.left_indent = Inches(0.4)
        p_cmd.paragraph_format.space_after = Pt(2)
        run_cmd = p_cmd.add_run(cmd)
        run_cmd.font.name = 'Consolas'
        run_cmd.font.size = Pt(9.5)
        if cmd.startswith('#'):
            run_cmd.font.color.rgb = RGBColor(128, 128, 128)
            run_cmd.font.italic = True
        else:
            run_cmd.font.color.rgb = RGBColor(0, 0, 0)
            
    doc.add_paragraph()

    # Section: Frontend Setup
    h5 = doc.add_paragraph()
    run_h5 = h5.add_run("5. Frontend Setup & Run Instructions")
    run_h5.font.bold = True
    run_h5.font.size = Pt(14)
    run_h5.font.color.rgb = RGBColor(74, 99, 49)
    
    doc.add_paragraph("Open a second, separate terminal/command line window and execute:")
    
    commands_fe = [
        "cd ttcell",
        "# Install dependencies",
        "npm install",
        "# Start frontend server",
        "npm run dev"
    ]
    for cmd in commands_fe:
        p_cmd = doc.add_paragraph()
        p_cmd.paragraph_format.left_indent = Inches(0.4)
        p_cmd.paragraph_format.space_after = Pt(2)
        run_cmd = p_cmd.add_run(cmd)
        run_cmd.font.name = 'Consolas'
        run_cmd.font.size = Pt(9.5)
        if cmd.startswith('#'):
            run_cmd.font.color.rgb = RGBColor(128, 128, 128)
            run_cmd.font.italic = True
        else:
            run_cmd.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("Open your browser and navigate to: http://localhost:3000")

    doc.add_paragraph()

    # Section: Credentials
    h6 = doc.add_paragraph()
    run_h6 = h6.add_run("6. Seed User Credentials")
    run_h6.font.bold = True
    run_h6.font.size = Pt(14)
    run_h6.font.color.rgb = RGBColor(74, 99, 49)
    
    p_cred = doc.add_paragraph("Once databases are seeded, log in with the following default accounts:")
    
    # Table of credentials
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Shading Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Role'
    hdr_cells[1].text = 'Email'
    hdr_cells[2].text = 'Default Password'
    
    row1 = table.rows[1].cells
    row1[0].text = 'Administrator'
    row1[1].text = 'admin@ttcell'
    row1[2].text = 'password'
    
    row2 = table.rows[2].cells
    row2[0].text = 'Trainee'
    row2[1].text = 'trainee@ttcell'
    row2[2].text = 'ChangeMeOnFirstLogin!'

    # Save
    out_path = os.path.join("d:\\TT-CELL-VOCATIONAL-TRAINING-main", "run guide.docx")
    doc.save(out_path)
    print(f"Document saved to {out_path}")

if __name__ == '__main__':
    create_guide()
